"""Shared pytest fixtures and helpers."""
import os
import tempfile

import pytest
from unittest.mock import Mock

from app import app
from fileExtraction import CONFIG

try:
    from docx import Document as DocxDocument
    DOCX_CREATE_AVAILABLE = True
except ImportError:
    DOCX_CREATE_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    PDF_CREATE_AVAILABLE = True
except ImportError:
    PDF_CREATE_AVAILABLE = False

try:
    from openpyxl import Workbook
    XLSX_CREATE_AVAILABLE = True
except ImportError:
    XLSX_CREATE_AVAILABLE = False


@pytest.fixture
def client():
    app.config["TESTING"] = True
    original_api_key = CONFIG.get("FILE_EXTRACTOR_KEY", "")
    CONFIG["FILE_EXTRACTOR_KEY"] = ""
    try:
        with app.test_client() as test_client:
            yield test_client
    finally:
        CONFIG["FILE_EXTRACTOR_KEY"] = original_api_key


def auth_headers():
    api_key = CONFIG.get("FILE_EXTRACTOR_KEY", "")
    if not api_key:
        return {}
    return {"Authorization": f"Bearer {api_key}"}


def configure_mock_download(mock_get, file_path, content_type="application/octet-stream"):
    with open(file_path, "rb") as handle:
        file_content = handle.read()
    mock_response = Mock()
    mock_response.iter_content = Mock(return_value=[file_content])
    mock_response.headers = {"Content-Type": content_type}
    mock_response.raise_for_status = Mock()
    mock_get.return_value = mock_response


@pytest.fixture
def sample_txt_file():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as handle:
        handle.write("Hello, this is a test text file.\nIt has multiple lines.\nFor testing purposes.")
        temp_path = handle.name
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)


@pytest.fixture
def sample_csv_file():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False, encoding="utf-8") as handle:
        handle.write("Name,Age,City\nJohn,30,New York\nJane,25,Los Angeles\nBob,35,Chicago")
        temp_path = handle.name
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)


@pytest.fixture
def sample_docx_file():
    if not DOCX_CREATE_AVAILABLE:
        pytest.skip("python-docx not available")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        temp_path = handle.name
    doc = DocxDocument()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It contains multiple paragraphs.")
    doc.add_paragraph("For testing the extraction functionality.")
    doc.save(temp_path)
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)


@pytest.fixture
def sample_pdf_file():
    if PDF_CREATE_AVAILABLE:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as handle:
            temp_path = handle.name
        pdf = canvas.Canvas(temp_path, pagesize=letter)
        pdf.drawString(100, 750, "Test PDF Content for Extraction")
        pdf.drawString(100, 730, "Second line of PDF text")
        pdf.save()
    else:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as handle:
            handle.write(
                b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> >>
endobj
4 0 obj
<< /Length 65 >>
stream
BT
/F1 12 Tf
100 700 Td
(Test PDF Content for Extraction) Tj
0 -20 Td
(Second line of PDF text) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000256 00000 n 
trailer
<< /Size 5 /Root 1 0 R >>
startxref
380
%%EOF"""
            )
            temp_path = handle.name
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)


@pytest.fixture
def sample_xlsx_file():
    if not XLSX_CREATE_AVAILABLE:
        pytest.skip("openpyxl not available")
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as handle:
        temp_path = handle.name
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Name", "Age", "City"])
    sheet.append(["Alice", 28, "Boston"])
    sheet.append(["Bob", 35, "Chicago"])
    workbook.save(temp_path)
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)
